from flask import send_file
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm
from flask import Flask, render_template, request
from app import app


@app.route('/')
def student():
    return render_template('student.html')


@app.route('/result', methods=['POST', 'GET'])
def result():
    if request.method == 'POST':
        result = request.form
        clg_name1 = request.form['clg name']
        print(clg_name1)
        department = request.form['department']
        dispach_date = request.form['dispach date']
        start_date = request.form['start date']
        end_date = request.form['end date']
        class_name = request.form['class name']
        df1 = pd.read_excel(request.files.get('file'))
        print("Loading....")
        #print(list(df1.columns[3]))
        sub_name = list(df1.columns)
        sub_1_name = sub_name[3]
        sub_2_name = sub_name[5]
        sub_3_name = sub_name[7]
        sub_4_name = sub_name[9]
        sub_5_name = sub_name[11]

        #print(sub_1_name)
        sub_1_UT1 = sub_name[4]
        sub_2_UT1 = sub_name[6]
        sub_3_UT1 = sub_name[8]
        sub_4_UT1 = sub_name[10]
        sub_5_UT1 = sub_name[12]
        maindata = []
        for i in range(0, len(df1)):
            data = {}
            total = 0
            data['Sr No'] = df1['Sr No'][i]
            data['Roll No'] = df1['Roll No'][i]
            data['Name of Students'] = df1['Name of Students'][i]
            # Reading Subject Attendance
            data['Sub1'] = round(df1[sub_1_name][i])
            data['Sub2'] = round(df1[sub_2_name][i])
            data['Sub3'] = round(df1[sub_3_name][i])
            data['Sub4'] = round(df1[sub_4_name][i])
            data['Sub5'] = round(df1[sub_5_name][i])
            data['Total'] = round(df1['Total'][i])
            # Reading UT1 marks
            data['Sub1_UT1'] = df1[sub_1_UT1][i]
            data['Sub2_UT1'] = df1[sub_2_UT1][i]
            data['Sub3_UT1'] = df1[sub_3_UT1][i]
            data['Sub4_UT1'] = df1[sub_4_UT1][i]
            data['Sub5_UT1'] = df1[sub_5_UT1][i]
            data['TG_Name'] = df1['TG_Name'][i]
            data['TG_Phone'] = df1['TG_Phone'][i]
            data['HOD'] = df1['HOD'][i]
            data['Principal'] = df1['Principal'][i]
            data['Parents_Name'] = df1['Parents Name'][i]
            data['Address'] = df1['Address'][i]
            data['Parents_Phone'] = df1['Parents Phone'][i]
            # print(data['Sub1-UT1'])
            maindata.append(data)
        # pprint.pprint(maindata)
        # -----Main Program----------------------
        Sr_No, Roll_No, Name_of_Students, Sub1, Sub2, Sub3, Sub4, Sub5, Total = [[] for bb in range(9)]
        Sub1_UT1, Sub2_UT1, Sub3_UT1, Sub4_UT1, Sub5_UT1 = [[] for aa in range(5)]
        TG, Phone, HOD1, Principal1, Parents1, Address1, ParentsPhone1 = [[] for cc in range(7)]
        for i in range(0, len(maindata)):
           Sr_No.append(maindata[i]['Sr No'])
           Roll_No.append(maindata[i]['Roll No'])
           Name_of_Students.append(maindata[i]['Name of Students'])
           Sub1.append(maindata[i]['Sub1'])
           Sub2.append(maindata[i]['Sub2'])
           Sub3.append(maindata[i]['Sub3'])
           Sub4.append(maindata[i]['Sub4'])
           Sub5.append(maindata[i]['Sub5'])
           Total.append(maindata[i]['Total'])

           Sub1_UT1.append(maindata[i]['Sub1_UT1'])
           Sub2_UT1.append(maindata[i]['Sub2_UT1'])
           Sub3_UT1.append(maindata[i]['Sub3_UT1'])
           Sub4_UT1.append(maindata[i]['Sub4_UT1'])
           Sub5_UT1.append(maindata[i]['Sub5_UT1'])

           TG.append(maindata[i]['TG_Name'])
           Phone.append(maindata[i]['TG_Phone'])
           HOD1.append(maindata[i]['HOD'])
           Principal1.append(maindata[i]['Principal'])

           Parents1.append(maindata[i]['Parents_Name'])
           Address1.append(maindata[i]['Address'])
           ParentsPhone1.append(maindata[i]['Parents_Phone'])
           # print(Principal1)
           # print("done")
        # -----Creating Word Document------------------------
        document = Document()
        sections = document.sections
        for section in sections:
           section.top_margin = Cm(1.0)
           section.bottom_margin = Cm(1.0)
           section.left_margin = Cm(1.75)
           section.right_margin = Cm(1.75)

        for (Sr, Roll, Name, S1, S2, S3, S4, S5, Total1, ST1, ST2, ST3, ST4, ST5, TG2, Phone2, HOD2, Principal2, Parents2, Address2, ParentsPhone2) in zip(
               Sr_No, Roll_No, Name_of_Students, Sub1, Sub2, Sub3, Sub4, Sub5, Total, Sub1_UT1, Sub2_UT1, Sub3_UT1,
               Sub4_UT1, Sub5_UT1, TG, Phone, HOD1, Principal1, Parents1, Address1, ParentsPhone1):

           clg_name = document.add_paragraph(clg_name1)
           clg_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
           dept_name = document.add_paragraph(department)
           dept_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
           date1 = document.add_paragraph('Date' + dispach_date)
           date1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
           ar1 = document.add_paragraph('ACADEMIC REPORT')
           ar1.alignment = WD_ALIGN_PARAGRAPH.CENTER
           period = document.add_paragraph('Period :' + start_date + ' to ' + end_date)
           period.alignment = WD_ALIGN_PARAGRAPH.CENTER
           sname = document.add_paragraph('Students Name : ')
           sname.add_run(str(Name)).bold = True
           roll = document.add_paragraph('Class : ' + class_name)
           roll.add_run('\t\t\t\t\t\t\tRoll No. : ' + str(Roll)).bold = True
           tg = document.add_paragraph()
           tg.add_run('Name of Teacher Guardian :' + TG2 + '\n Teacher Guardian Phone :' + str(Phone2)).bold = True

           ar = document.add_paragraph('ATTENDANCE RECORD')
           ar.alignment = WD_ALIGN_PARAGRAPH.CENTER
           overallatt = document.add_paragraph()
           overallatt.add_run('\t\t\t\t\tTotal Attendance : ' + str(Total1) + '%').bold = True
           # --------table---------------
           table = document.add_table(rows=6, cols=5, style='Table Grid')
           table.cell(0, 0).width = 731520  # 0.8 * 914400
           table.cell(1, 0).width = 731520
           table.cell(2, 0).width = 731520
           table.cell(3, 0).width = 731520
           table.cell(4, 0).width = 731520  # 1.2 * 914400
           table.cell(5, 0).width = 731520

           table.cell(0, 1).width = 4389120  # 4.8 * 914400
           table.cell(1, 1).width = 4389120
           table.cell(2, 1).width = 4389120
           table.cell(3, 1).width = 4389120
           table.cell(4, 1).width = 4389120  # 1.2 * 914400
           table.cell(5, 1).width = 4389120

           table.cell(0, 2).width = 1097280  # 1.2 * 914400
           table.cell(1, 2).width = 1097280
           table.cell(2, 2).width = 1097280
           table.cell(3, 2).width = 1097280
           table.cell(4, 2).width = 1097280  # 1.2 * 914400
           table.cell(5, 2).width = 1097280

           table.cell(0, 3).width = 1828800  # 1.2 * 914400
           table.cell(1, 3).width = 1828800
           table.cell(2, 3).width = 1828800
           table.cell(3, 3).width = 1828800
           table.cell(4, 3).width = 1828800  # 1.5 * 914400
           table.cell(5, 3).width = 1828800

           table.cell(0, 4).width = 1097280  # 1.2 * 914400
           table.cell(1, 4).width = 1097280
           table.cell(2, 4).width = 1097280
           table.cell(3, 4).width = 1097280
           table.cell(4, 4).width = 1097280  # 1.2 * 914400
           table.cell(5, 4).width = 1097280

           hdr_cells = table.rows[0].cells
           hdr_cells[0].text = 'Sr. No.'
           hdr_cells[1].text = 'Subject Name'
           hdr_cells[2].text = 'Attendance in %'
           hdr_cells[3].text = 'Unit Test 1 Marks (out of 25)'
           hdr_cells[4].text = 'Remark'


           hdr_cells = table.rows[1].cells
           hdr_cells[0].text = '1'
           hdr_cells[1].text = sub_1_name
           hdr_cells[2].text = str(S1)
           hdr_cells[3].text = str(ST1)
           if str(ST1) == 'AB':
               hdr_cells[4].text = 'Absent'
           elif str(ST1) == '':
               hdr_cells[4].text = 'Not Applicable'

           else:
               if ST1 >= 10:
                   hdr_cells[4].text = 'Pass'
               else:
                   hdr_cells[4].text = 'Fail'

           hdr_cells = table.rows[2].cells
           hdr_cells[0].text = '2'
           hdr_cells[1].text = sub_2_name
           hdr_cells[2].text = str(S2)
           hdr_cells[3].text = str(ST2)
           if str(ST2) == 'AB':
               hdr_cells[4].text = 'Absent'
           else:
               if ST2 >= 10:
                   hdr_cells[4].text = 'Pass'
               else:
                   hdr_cells[4].text = 'Fail'
           hdr_cells = table.rows[3].cells
           hdr_cells[0].text = '3'
           hdr_cells[1].text = sub_3_name
           hdr_cells[2].text = str(S3)
           hdr_cells[3].text = str(ST3)
           if str(ST3) == 'AB':
               hdr_cells[4].text = 'Absent'
           else:
               if ST3 >= 10:
                   hdr_cells[4].text = 'Pass'
               else:
                   hdr_cells[4].text = 'Fail'
           hdr_cells = table.rows[4].cells
           hdr_cells[0].text = '4'
           hdr_cells[1].text = sub_4_name
           hdr_cells[2].text = str(S4)
           hdr_cells[3].text = str(ST4)
           if str(ST4) == 'AB':
               hdr_cells[4].text = 'Absent'
           else:
               if ST4 >= 10:
                   hdr_cells[4].text = 'Pass'
               else:
                   hdr_cells[4].text = 'Fail'

           hdr_cells = table.rows[5].cells
           hdr_cells[0].text = '5'
           hdr_cells[1].text = sub_5_name
           hdr_cells[2].text = str(S5)
           hdr_cells[3].text = str(ST5)
           if str(ST5) == 'AB':
               hdr_cells[4].text = 'Absent'
           else:
               if ST5 >= 10:
                   hdr_cells[4].text = 'Pass'
               else:
                   hdr_cells[4].text = 'Fail'
           remark = document.add_paragraph('\nREMARK')
           remark.alignment = WD_ALIGN_PARAGRAPH.CENTER
           rm = document.add_paragraph()
           rm.add_run(
               '1) Your ward is not attending lectures.\n2)  Performance of ward in attendance is very poor. \n3)  Attendance of your ward is below 75%; his term is liable to be detained.')

           rm1 = document.add_paragraph()
           rm1.add_run(
               'In case of any queries about the performance of your ward, Please visit   personally to the college or contact the concerned teacher guardian.\n\n').bold = True

           table1 = document.add_table(rows=2, cols=3)
           hdr_cells = table1.rows[0].cells
           hdr_cells[0].text = TG2
           hdr_cells[1].text = HOD2
           hdr_cells[2].text = Principal2

           hdr_cells = table1.rows[1].cells
           hdr_cells[0].text = 'Teacher Guradian'
           hdr_cells[1].text = '  HOD'
           hdr_cells[2].text = '  Principal\n\n\n'
           pname = document.add_paragraph()
           pname.add_run(' \tTo \n \t' + Parents2 + '\n\t' + Address2 + '\n\t' + 'Mobile No.:' + str(ParentsPhone2)).bold = True
           document.add_page_break()
        document.save('Parents_Letter.docx')
        print("well done, Letters generated")
        return render_template("result.html",result = result, department=department, start_date=start_date,end_date=end_date,
                              class_name=class_name,dispach_date=dispach_date)

@app.route('/sample')
def downloadSample ():
    #For windows you need to use drive name [ex: F:/Example.pdf]
    path = r"F:\Letter-Writing-soft\flask_letter_writing_soft\sample.xlsx"
    return send_file(path, as_attachment=True)

@app.route('/download')
def downloadFile ():
    #For windows you need to use drive name [ex: F:/Example.pdf]
    path = r"F:\Letter-Writing-soft\flask_letter_writing_soft\Parents_Letter.docx"
    return send_file(path, as_attachment=True)

@app.route('/about')
def about():
    return render_template("about.html")

if __name__ == '__main__':
   app.run(debug = True)